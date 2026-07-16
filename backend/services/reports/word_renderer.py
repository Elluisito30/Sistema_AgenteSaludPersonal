"""
Word Renderer — Generacion de reportes Word con python-docx.

Solo renderiza. No accede a BD, no accede al modelo ML, no recalcula.
Consume exclusivamente report_data (dict de ReportEngine).
"""

import io
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from services.reports.styles import get_translations, get_risk_colors, COLORS
from services.reports.helpers import (
    format_risk_badge,
    format_confidence_badge,
    format_bmi_indicator,
    format_weight_projection,
    format_macronutrients,
    format_exercise_plan,
    format_profile_summary,
)


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            table.rows[r + 1].cells[c].text = str(val)
    return table


# ==========================================
# HEALTH REPORT WORD
# ==========================================

def generate_health_word(report_data, plot_bytes=None, language="es"):
    t = get_translations(language)
    doc = Document()
    profile_info = format_profile_summary(report_data.get("profile"))
    analysis = report_data.get("analysis", {})
    risk_badge = format_risk_badge(analysis.get("health_risk"), language)
    bmi_ind = format_bmi_indicator(
        analysis.get("bmi"), analysis.get("bmi_category"), language
    )
    c_rgb = risk_badge["color_rgb"]

    # Titulo
    title = doc.add_heading(t["clinical_risk_report"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run(f"{t['date_generated']}: {datetime.now().strftime('%d/%m/%Y')}\n").bold = True
    p.add_run(t["patient_evaluation"])

    # 1. Evaluacion Clinica
    doc.add_heading(t["current_clinical_evaluation"], level=1)
    p_risk = doc.add_paragraph(t["calculated_risk"] + " ")
    r = p_risk.add_run(risk_badge["label"])
    r.bold = True
    r.font.color.rgb = RGBColor(*c_rgb)

    doc.add_paragraph(f"{t['health_score']}: {analysis.get('health_score', 0)} / 100")
    doc.add_paragraph(f"{t['bmi']}: {bmi_ind['value']} ({bmi_ind['label']})")
    doc.add_paragraph(f"{t['bmr']}: {analysis.get('bmr', 0)}")
    doc.add_paragraph(f"{t['tdee']}: {analysis.get('tdee', 0)}")

    # 2. Factores
    doc.add_heading(t["factors_and_comorbidities"], level=1)
    profile = report_data.get("profile", {})
    chronic = profile.get("chronic_diseases", [])
    genetic = profile.get("genetic_risk_factors", [])
    _add_table(doc,
        [t["category"], t["detected_conditions"]],
        [
            [t["chronic"], ", ".join(chronic) if chronic else t["no_record"]],
            [t["genetic_factors"], ", ".join(genetic) if genetic else t["no_record"]],
        ],
    )

    # 3. Alertas
    alerts_data = report_data.get("alerts", {})
    all_alerts = alerts_data.get("high", []) + alerts_data.get("medium", [])
    if all_alerts:
        doc.add_heading(t["alerts"], level=1)
        rows = [[al.get("type", ""), al.get("message", ""), al.get("priority", "")] for al in all_alerts[:6]]
        _add_table(doc, [t["type"], t["message"], t["value"]], rows)

    # 4. Objetivos
    goals = report_data.get("goals", {}).get("all", [])
    if goals:
        doc.add_heading(t["objectives"], level=1)
        for g in goals[:5]:
            doc.add_paragraph(g, style="List Bullet")

    # 5. Narrativas
    narratives = report_data.get("narratives", {})
    for nk in ["clinical_summary", "bmi_interpretation", "risk_interpretation"]:
        nav = narratives.get(nk)
        if nav and nav.get("text"):
            doc.add_heading(nav["title"], level=1)
            doc.add_paragraph(nav["text"])

    # 6. ML + XAI
    ml = report_data.get("ml_prediction", {})
    if ml.get("available"):
        doc.add_heading(t["ml_prediction"], level=1)
        conf = format_confidence_badge(ml.get("confidence", 0), language)
        _add_table(doc,
            [t["metric"], t["value"]],
            [
                [t["classification"], ml.get("predicted_class_label", "")],
                [t["confidence"], conf["label"]],
                [t["model"], ml.get("model_used", "")],
            ],
        )

    xai = report_data.get("xai", {})
    if xai.get("available"):
        doc.add_heading(t["xai_explanation"], level=1)
        if xai.get("summary"):
            doc.add_paragraph(xai["summary"])
        if xai.get("main_reason"):
            doc.add_paragraph(xai["main_reason"])

    # 7. Grafico
    if plot_bytes:
        doc.add_heading(t["graph_projection"], level=1)
        doc.add_picture(io.BytesIO(plot_bytes), width=Inches(6.0))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ==========================================
# PREDICTIONS REPORT WORD
# ==========================================

def generate_predictions_word(report_data, language="es"):
    t = get_translations(language)
    doc = Document()
    profile_info = format_profile_summary(report_data.get("profile"))

    title = doc.add_heading(t["predictions_report"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run(f"{t['date_generated']}: {datetime.now().strftime('%d/%m/%Y')}\n").bold = True
    p.add_run(t["patient_evaluation"])

    # Proyecciones
    doc.add_heading(t["weight_summary"], level=1)
    wp = format_weight_projection(report_data.get("predictions"))
    if wp.get("available"):
        rows = [[p["label"], f"{p['weight_kg']} kg"] for p in wp["periods"]]
        _add_table(doc, [t["period"], t["weight"]], rows)
    else:
        doc.add_paragraph(t["no_projection_data"])

    # ML
    ml = report_data.get("ml_prediction", {})
    if ml.get("available"):
        doc.add_heading(t["ml_prediction"], level=1)
        conf = format_confidence_badge(ml.get("confidence", 0), language)
        _add_table(doc,
            [t["metric"], t["value"]],
            [
                [t["classification"], ml.get("predicted_class_label", "")],
                [t["confidence"], conf["label"]],
                [t["model"], ml.get("model_used", "")],
            ],
        )

    # Narrativas
    narratives = report_data.get("narratives", {})
    for nk in ["ml_interpretation", "prognosis"]:
        nav = narratives.get(nk)
        if nav and nav.get("text"):
            doc.add_heading(nav["title"], level=1)
            doc.add_paragraph(nav["text"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ==========================================
# NUTRITION REPORT WORD
# ==========================================

def generate_nutrition_word(report_data, language="es"):
    t = get_translations(language)
    doc = Document()

    title = doc.add_heading(t["recipes_report"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run(f"{t['date_generated']}: {datetime.now().strftime('%d/%m/%Y')}\n").bold = True
    p.add_run(t["patient_evaluation"])

    # Resumen nutricional
    doc.add_heading(t["nutrition_summary_1"], level=1)
    macros = format_macronutrients(report_data.get("nutrition"))
    if macros.get("available"):
        doc.add_paragraph(f"{t['daily_calories']}: {macros['daily_calories']} kcal")
        doc.add_paragraph(f"{t['protein']}: {macros['protein_g']} g ({macros['protein_pct']}%)")
        doc.add_paragraph(f"{t['carbs']}: {macros['carbs_g']} g ({macros['carbs_pct']}%)")
        doc.add_paragraph(f"{t['fats']}: {macros['fats_g']} g ({macros['fats_pct']}%)")

        doc.add_heading(t["suggested_recipes"], level=1)
        protein = macros["protein_g"]
        rows = [
            [t["breakfast"], "Avena con frutas, chia y nueces" if language == "es" else "Oatmeal with fruits, chia and nuts", str(round(protein * 0.2))],
            [t["lunch"], "Pechuga de pollo con quinoa" if language == "es" else "Grilled chicken with quinoa", str(round(protein * 0.4))],
            [t["dinner"], "Ensalada de salmon" if language == "es" else "Salmon salad", str(round(protein * 0.3))],
            [t["snack"], "Yogur griego o almendras" if language == "es" else "Greek yogurt or almonds", str(round(protein * 0.1))],
        ]
        _add_table(doc, [t["meal"], t["recipe"], t["protein_g"]], rows)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ==========================================
# EXERCISE REPORT WORD
# ==========================================

def generate_exercise_word(report_data, language="es"):
    t = get_translations(language)
    doc = Document()

    title = doc.add_heading(t["exercise_report"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run(f"{t['date_generated']}: {datetime.now().strftime('%d/%m/%Y')}\n").bold = True
    p.add_run(t["patient_evaluation"])

    doc.add_heading(t["exercise_plan_1"], level=1)
    ex = format_exercise_plan(report_data.get("exercise"))
    if ex.get("available"):
        _add_table(doc,
            [t["type"], t["duration"]],
            [
                [t["cardio"], ex["cardio"]],
                [t["strength"], ex["strength"]],
                [t["flexibility"], ex["flexibility"]],
            ],
        )
    else:
        doc.add_paragraph(t["no_data"])

    # Narrativas
    narratives = report_data.get("narratives", {})
    for nk in ["risk_interpretation", "recommendations_summary"]:
        nav = narratives.get(nk)
        if nav and nav.get("text"):
            doc.add_heading(nav["title"], level=1)
            doc.add_paragraph(nav["text"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
