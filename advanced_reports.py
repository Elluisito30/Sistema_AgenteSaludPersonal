import io
import pandas as pd
import xlsxwriter
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from datetime import datetime
import plotly.graph_objects as go
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# TRADUCCIONES
# ==========================================
TRANSLATIONS = {
    "es": {
        "app_name": "🏥 HEALTH AI",
        "health_report": "Informe Clínico",
        "predictions_report": "Predicciones de Peso",
        "recipes_report": "Plan de Alimentación",
        "exercise_report": "Plan de Ejercicio",
        "date_generated": "Fecha de Generación",
        "patient": "Paciente",
        "health_score": "Puntuación de Salud",
        "risk_level": "Nivel de Riesgo",
        "bmi": "IMC",
        "nutrition_summary": "Resumen Nutricional",
        "daily_calories": "Calorías Diarias",
        "protein": "Proteínas",
        "meal": "Momento",
        "recipe": "Receta Sugerida",
        "protein_g": "Proteínas (g)",
        "breakfast": "Desayuno",
        "lunch": "Almuerzo",
        "dinner": "Cena",
        "snack": "Snack",
        "exercise_plan": "Plan de Ejercicio Semanal",
        "cardio": "Cardio",
        "strength": "Fuerza",
        "flexibility": "Flexibilidad",
        "current_weight": "Peso Actual",
        "2_weeks": "2 Semanas",
        "1_month": "1 Mes",
        "6_months": "6 Meses",
        "weight": "Peso (kg)",
        "metric": "Métrica",
        "value": "Valor",
        "clinical_factors": "Factores de Riesgo Documentados",
        "type": "Tipo",
        "conditions": "Condiciones (Carga Negativa)",
        "chronic": "Enfermedades Crónicas",
        "genetic": "Riesgo Genético Familiar",
        "no_data": "Ninguna registrada",
        "history": "Historial",
        "timeline": "Evolución del Riesgo (Histórico)",
        "indicator": "Indicador",
        "category": "Categoría",
        "detected_conditions": "Condiciones Detectadas",
        "genetics": "Genética",
        "risk_evaluation": "Evaluación de Riesgo Actual",
        "clinical_status": "Estado Clínico",
        "clinical_risk_report": "Informe de Riesgo Cardiovascular y Metabólico",
        "current_clinical_evaluation": "1. Evaluación Clínica Actual",
        "calculated_risk": "Riesgo Calculado:",
        "factors_and_comorbidities": "2. Factores y Comorbilidades",
        "genetic_factors": "Factores Genéticos",
        "no_record": "Sin registro",
        "graph_projection": "3. Proyección Gráfica",
        "weight_summary": "Resumen de Peso",
        "period": "Periodo",
        "nutrition_summary_1": "1. Resumen Nutricional",
        "suggested_recipes": "2. Recetas Sugeridas",
        "exercise_plan_1": "1. Plan de Ejercicio Semanal",
        "patient_evaluation": "Evaluación del Paciente: Confidencial",
    },
    "en": {
        "app_name": "🏥 HEALTH AI",
        "health_report": "Clinical Report",
        "predictions_report": "Weight Predictions",
        "recipes_report": "Nutrition Plan",
        "exercise_report": "Exercise Plan",
        "date_generated": "Generated Date",
        "patient": "Patient",
        "health_score": "Health Score",
        "risk_level": "Risk Level",
        "bmi": "BMI",
        "nutrition_summary": "Nutrition Summary",
        "daily_calories": "Daily Calories",
        "protein": "Protein",
        "meal": "Meal",
        "recipe": "Suggested Recipe",
        "protein_g": "Protein (g)",
        "breakfast": "Breakfast",
        "lunch": "Lunch",
        "dinner": "Dinner",
        "snack": "Snack",
        "exercise_plan": "Weekly Exercise Plan",
        "cardio": "Cardio",
        "strength": "Strength",
        "flexibility": "Flexibility",
        "current_weight": "Current Weight",
        "2_weeks": "2 Weeks",
        "1_month": "1 Month",
        "6_months": "6 Months",
        "weight": "Weight (kg)",
        "metric": "Metric",
        "value": "Value",
        "clinical_factors": "Documented Risk Factors",
        "type": "Type",
        "conditions": "Conditions (Negative Load)",
        "chronic": "Chronic Diseases",
        "genetic": "Family Genetic Risk",
        "no_data": "None recorded",
        "history": "History",
        "timeline": "Risk Evolution (Historical)",
        "indicator": "Indicator",
        "category": "Category",
        "detected_conditions": "Detected Conditions",
        "genetics": "Genetics",
        "risk_evaluation": "Current Risk Evaluation",
        "clinical_status": "Clinical Status",
        "clinical_risk_report": "Cardiovascular and Metabolic Risk Report",
        "current_clinical_evaluation": "1. Current Clinical Evaluation",
        "calculated_risk": "Calculated Risk:",
        "factors_and_comorbidities": "2. Factors and Comorbidities",
        "genetic_factors": "Genetic Factors",
        "no_record": "No record",
        "graph_projection": "3. Graph Projection",
        "weight_summary": "Weight Summary",
        "period": "Period",
        "nutrition_summary_1": "1. Nutrition Summary",
        "suggested_recipes": "2. Suggested Recipes",
        "exercise_plan_1": "1. Weekly Exercise Plan",
        "patient_evaluation": "Patient Evaluation: Confidential",
    }
}

def get_translations(language="es"):
    """Get translation dictionary for given language, default to Spanish"""
    return TRANSLATIONS.get(language, TRANSLATIONS["es"])

# ==========================================
# UTILIDADES DE COLOR Y RIESGO
# ==========================================
def get_risk_colors(risk_level):
    """Retorna colores según el riesgo (Hex para Excel, Color para PDF)"""
    if "Crítico" in risk_level or "Critical" in risk_level:
        return {"hex": "#e74c3c", "pdf": colors.HexColor("#e74c3c"), "rgb": (231, 76, 60), "name": "Crítico"}
    elif "En Riesgo" in risk_level or "At Risk" in risk_level:
        return {"hex": "#e67e22", "pdf": colors.HexColor("#e67e22"), "rgb": (230, 126, 34), "name": "En Riesgo"}
    elif "Aceptable" in risk_level or "Acceptable" in risk_level:
        return {"hex": "#f1c40f", "pdf": colors.HexColor("#f1c40f"), "rgb": (241, 196, 15), "name": "Aceptable"}
    elif "Bueno" in risk_level or "Good" in risk_level:
        return {"hex": "#2ecc71", "pdf": colors.HexColor("#2ecc71"), "rgb": (46, 204, 113), "name": "Bueno"}
    else:
        return {"hex": "#7f8c8d", "pdf": colors.HexColor("#7f8c8d"), "rgb": (127, 140, 141), "name": "Sin Datos"}

# ==========================================
# 1. EXPORTACIÓN A PDF (ReportLab)
# ==========================================
def generate_risk_report_pdf(analysis_data, profile_data, score_history, plot_bytes, language="es"):
    """Genera un PDF profesional con ReportLab usando Flowables"""
    t = get_translations(language)
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'MainTitle', parent=styles['Heading1'], fontSize=18, spaceAfter=20, textColor=colors.HexColor("#2c3e50"), alignment=1
    )
    subtitle_style = ParagraphStyle('SubTitle', parent=styles['Heading2'], fontSize=14, spaceAfter=10, textColor=colors.HexColor("#34495e"))
    normal_style = styles["Normal"]
    
    elements = []
    
    # 1. ENCABEZADO
    elements.append(Paragraph(f"{t['app_name']} - {t['health_report']}", title_style))
    elements.append(Paragraph(f"<b>{t['date_generated']}:</b> {datetime.now().strftime('%Y-%m-%d %H:%M')}", normal_style))
    elements.append(Paragraph(f"<b>{t['patient']}:</b> {profile_data.get('email', 'Usuario')}", normal_style))
    elements.append(Spacer(1, 20))
    
    risk_level = analysis_data.get("health_risk", t['no_data'])
    c_info = get_risk_colors(risk_level)
    
    # 2. SECCIÓN: ESTADO ACTUAL
    elements.append(Paragraph(t['risk_evaluation'], subtitle_style))
    
    # Tabla de métricas principales
    metrics_data = [
        [t['metric'], t['value']],
        [t['health_score'], f"{analysis_data.get('health_score', 0)} / 100"],
        [t['clinical_status'], risk_level],
        [t['bmi'], str(analysis_data.get('bmi', 0))]
    ]
    
    t_metrics = Table(metrics_data, colWidths=[200, 200])
    t_metrics.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (1, 0), colors.HexColor("#34495e")),
        ('TEXTCOLOR', (0, 0), (1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor("#ecf0f1")),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        # Colorear la fila de Estado Clínico condicionalmente
        ('BACKGROUND', (1, 2), (1, 2), c_info["pdf"]),
        ('TEXTCOLOR', (1, 2), (1, 2), colors.white)
    ]))
    elements.append(t_metrics)
    elements.append(Spacer(1, 20))
    
    # 3. SECCIÓN: FACTORES CLÍNICOS
    elements.append(Paragraph(t['clinical_factors'], subtitle_style))
    chronic = profile_data.get("chronic_diseases", [])
    genetic = profile_data.get("genetic_risk_factors", [])
    
    factors_data = [[t['type'], t['conditions']]]
    factors_data.append([t['chronic'], ", ".join(chronic) if chronic else t['no_data']])
    factors_data.append([t['genetic'], ", ".join(genetic) if genetic else t['no_data']])
    
    t_factors = Table(factors_data, colWidths=[150, 350])
    t_factors.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (1, 0), colors.HexColor("#2c3e50")),
        ('TEXTCOLOR', (0, 0), (1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
    ]))
    elements.append(t_factors)
    elements.append(Spacer(1, 20))
    
    # 4. GRÁFICO (Kaleido High Res Image)
    if plot_bytes:
        elements.append(Paragraph(t['timeline'], subtitle_style))
        img = Image(io.BytesIO(plot_bytes))
        # Escalar imagen para que encaje
        img.drawWidth = 6 * inch
        img.drawHeight = 3.5 * inch
        elements.append(img)
    
    doc.build(elements)
    buffer.seek(0)
    return buffer

# ==========================================
# 2. EXPORTACIÓN A EXCEL (XlsxWriter)
# ==========================================
def generate_risk_report_excel(analysis_data, profile_data, history_data, language="es"):
    """Genera Excel de 3 hojas con Formato Condicional (Sin Imágenes)"""
    t = get_translations(language)
    buffer = io.BytesIO()
    workbook = xlsxwriter.Workbook(buffer, {'in_memory': True})
    
    # FORMATOS
    header_format = workbook.add_format({'bold': True, 'bg_color': '#2c3e50', 'font_color': 'white', 'border': 1})
    
    # Formatos Condicionales para Riesgo
    format_critico = workbook.add_format({'bg_color': '#ffc7ce', 'font_color': '#9c0006'})
    format_riesgo = workbook.add_format({'bg_color': '#ffeb9c', 'font_color': '#9c6500'})
    format_bueno = workbook.add_format({'bg_color': '#c6efce', 'font_color': '#006100'})
    
    # HOJA 1: MÉTRICAS
    ws1 = workbook.add_worksheet(t['metrics'] if 'metrics' in t else 'Métricas')
    ws1.set_column('A:B', 25)
    ws1.write(0, 0, t['indicator'], header_format)
    ws1.write(0, 1, t['value'], header_format)
    
    metrics = [
        (t['health_score'], analysis_data.get('health_score', 0)),
        (t['risk_level'], analysis_data.get('health_risk', t['no_data'])),
        (t['bmi'], analysis_data.get('bmi', 0)),
        ('TMB', analysis_data.get('bmr', 0))
    ]
    
    for row_num, (ind, val) in enumerate(metrics, 1):
        ws1.write(row_num, 0, ind)
        ws1.write(row_num, 1, val)
        
    # Aplicar formato condicional a la celda del estado
    ws1.conditional_format('B3:B3', {'type': 'cell', 'criteria': '==', 'value': '"Crítico (Requiere atención médica urgente)"', 'format': format_critico})
    ws1.conditional_format('B3:B3', {'type': 'cell', 'criteria': '==', 'value': '"En Riesgo"', 'format': format_riesgo})
    ws1.conditional_format('B3:B3', {'type': 'cell', 'criteria': '==', 'value': '"Bueno"', 'format': format_bueno})
    
    # HOJA 2: HISTORIAL
    ws2 = workbook.add_worksheet(t['history'])
    ws2.set_column('A:C', 20)
    ws2.write_row(0, 0, ['Fecha', t['health_score'], t['risk_level']], header_format)
    
    for row_num, hist in enumerate(history_data, 1):
        ws2.write(row_num, 0, hist.get('analyzed_at', '')[:10])
        score = hist.get('health_score', 0)
        ws2.write(row_num, 1, score)
        ws2.write(row_num, 2, hist.get('health_risk', ''))
        
    # Formato condicional para puntuaciones en la Hoja 2
    ws2.conditional_format(f'B2:B{len(history_data)+1}', {'type': 'cell', 'criteria': '<=', 'value': 40, 'format': format_critico})
    ws2.conditional_format(f'B2:B{len(history_data)+1}', {'type': 'cell', 'criteria': '>=', 'value': 81, 'format': format_bueno})
    
    # HOJA 3: FACTORES DE RIESGO
    ws3 = workbook.add_worksheet(t['clinical_factors'])
    ws3.set_column('A:B', 30)
    ws3.write_row(0, 0, [t['category'], t['detected_conditions']], header_format)
    
    chronic = profile_data.get("chronic_diseases", [])
    genetic = profile_data.get("genetic_risk_factors", [])
    
    ws3.write(1, 0, t['chronic'])
    ws3.write(1, 1, ", ".join(chronic) if chronic else t['no_data'])
    
    ws3.write(2, 0, t['genetics'])
    ws3.write(2, 1, ", ".join(genetic) if genetic else t['no_data'])
    
    workbook.close()
    buffer.seek(0)
    return buffer

# ==========================================
# 3. EXPORTACIÓN A WORD (python-docx)
# ==========================================
def generate_risk_report_word(analysis_data, profile_data, plot_bytes, language="es"):
    """Genera un documento Word para entornos académicos o clínicos estandarizados"""
    t = get_translations(language)
    doc = Document()
    
    # Título Principal
    title = doc.add_heading(t['clinical_risk_report'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo y Fecha
    p = doc.add_paragraph()
    p.add_run(f"{t['date_generated']}: {datetime.now().strftime('%d/%m/%Y')}\n").bold = True
    p.add_run(t['patient_evaluation'])
    
    # Estado Clínico con Colores Nativos de Word
    doc.add_heading(t['current_clinical_evaluation'], level=1)
    risk = analysis_data.get('health_risk', t['no_data'])
    c_info = get_risk_colors(risk)
    
    p_risk = doc.add_paragraph(t['calculated_risk'] + ' ')
    r = p_risk.add_run(risk)
    r.bold = True
    r.font.color.rgb = RGBColor(*c_info["rgb"])
    
    p_score = doc.add_paragraph(f"{t['health_score']}: {analysis_data.get('health_score', 0)} / 100")
    
    # Tabla Clínica
    doc.add_heading(t['factors_and_comorbidities'], level=1)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = t['category']
    hdr_cells[1].text = t['detected_conditions']
    
    chronic = profile_data.get("chronic_diseases", [])
    genetic = profile_data.get("genetic_risk_factors", [])
    
    row_cells1 = table.rows[1].cells
    row_cells1[0].text = t['chronic']
    row_cells1[1].text = ", ".join(chronic) if chronic else t['no_record']
    
    row_cells2 = table.rows[2].cells
    row_cells2[0].text = t['genetic_factors']
    row_cells2[1].text = ", ".join(genetic) if genetic else t['no_record']
    
    # Gráfico
    if plot_bytes:
        doc.add_heading(t['graph_projection'], level=1)
        image_stream = io.BytesIO(plot_bytes)
        doc.add_picture(image_stream, width=Inches(6.0))
        
    # Guardar en buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 4. REPORTES DE PREDICCIONES
# ==========================================
def generate_predictions_report_pdf(analysis_data, profile_data, latest_prediction, language="es"):
    t = get_translations(language)
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('MainTitle', parent=styles['Heading1'], fontSize=18, spaceAfter=20, textColor=colors.HexColor("#2c3e50"), alignment=1)
    subtitle_style = ParagraphStyle('SubTitle', parent=styles['Heading2'], fontSize=14, spaceAfter=10, textColor=colors.HexColor("#34495e"))
    normal_style = styles["Normal"]

    elements = []
    elements.append(Paragraph(f"{t['app_name']} - {t['predictions_report']}", title_style))
    elements.append(Paragraph(f"<b>{t['date_generated']}:</b> {datetime.now().strftime('%Y-%m-%d %H:%M')}", normal_style))
    elements.append(Paragraph(f"<b>{t['patient']}:</b> {profile_data.get('email', 'Usuario')}", normal_style))
    elements.append(Spacer(1, 20))

    elements.append(Paragraph(t['weight_summary'], subtitle_style))
    weight_initial = latest_prediction.get('profile_snapshot', {}).get('weight', 0) if latest_prediction else 0
    weight_2w = latest_prediction.get('predictions_data', {}).get('predictions', {}).get('2_weeks', {}).get('weight_kg', 0) if latest_prediction else 0
    weight_1m = latest_prediction.get('predictions_data', {}).get('predictions', {}).get('1_month', {}).get('weight_kg', 0) if latest_prediction else 0
    weight_6m = latest_prediction.get('predictions_data', {}).get('predictions', {}).get('6_months', {}).get('weight_kg', 0) if latest_prediction else 0

    predictions_data = [
        [t['period'], t['weight']],
        [t['current_weight'], f"{weight_initial}"],
        [t['2_weeks'], f"{weight_2w}"],
        [t['1_month'], f"{weight_1m}"],
        [t['6_months'], f"{weight_6m}"]
    ]
    t_predictions = Table(predictions_data, colWidths=[200, 200])
    t_predictions.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (1, 0), colors.HexColor("#34495e")),
        ('TEXTCOLOR', (0, 0), (1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor("#ecf0f1")),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    elements.append(t_predictions)
    elements.append(Spacer(1, 20))

    doc.build(elements)
    buffer.seek(0)
    return buffer


def generate_predictions_report_excel(analysis_data, profile_data, latest_prediction, language="es"):
    t = get_translations(language)
    buffer = io.BytesIO()
    workbook = xlsxwriter.Workbook(buffer, {'in_memory': True})
    header_format = workbook.add_format({'bold': True, 'bg_color': '#2c3e50', 'font_color': 'white', 'border': 1})

    ws1 = workbook.add_worksheet(t['predictions_report'])
    ws1.set_column('A:B', 25)
    ws1.write(0, 0, t['period'], header_format)
    ws1.write(0, 1, t['weight'], header_format)

    weight_initial = latest_prediction.get('profile_snapshot', {}).get('weight', 0) if latest_prediction else 0
    weight_2w = latest_prediction.get('predictions_data', {}).get('predictions', {}).get('2_weeks', {}).get('weight_kg', 0) if latest_prediction else 0
    weight_1m = latest_prediction.get('predictions_data', {}).get('predictions', {}).get('1_month', {}).get('weight_kg', 0) if latest_prediction else 0
    weight_6m = latest_prediction.get('predictions_data', {}).get('predictions', {}).get('6_months', {}).get('weight_kg', 0) if latest_prediction else 0

    predictions = [
        (t['current_weight'], weight_initial),
        (t['2_weeks'], weight_2w),
        (t['1_month'], weight_1m),
        (t['6_months'], weight_6m)
    ]
    for row_num, (periodo, peso) in enumerate(predictions, 1):
        ws1.write(row_num, 0, periodo)
        ws1.write(row_num, 1, peso)

    workbook.close()
    buffer.seek(0)
    return buffer


def generate_predictions_report_word(analysis_data, profile_data, latest_prediction, language="es"):
    t = get_translations(language)
    doc = Document()
    title = doc.add_heading(t['predictions_report'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.add_run(f"{t['date_generated']}: {datetime.now().strftime('%d/%m/%Y')}\n").bold = True
    p.add_run(t['patient_evaluation'])

    doc.add_heading(t['weight_summary'], level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = t['period']
    hdr_cells[1].text = t['weight']

    weight_initial = latest_prediction.get('profile_snapshot', {}).get('weight', 0) if latest_prediction else 0
    weight_2w = latest_prediction.get('predictions_data', {}).get('predictions', {}).get('2_weeks', {}).get('weight_kg', 0) if latest_prediction else 0
    weight_1m = latest_prediction.get('predictions_data', {}).get('predictions', {}).get('1_month', {}).get('weight_kg', 0) if latest_prediction else 0
    weight_6m = latest_prediction.get('predictions_data', {}).get('predictions', {}).get('6_months', {}).get('weight_kg', 0) if latest_prediction else 0

    table.rows[1].cells[0].text = t['current_weight']
    table.rows[1].cells[1].text = str(weight_initial)
    table.rows[2].cells[0].text = t['2_weeks']
    table.rows[2].cells[1].text = str(weight_2w)
    table.rows[3].cells[0].text = t['1_month']
    table.rows[3].cells[1].text = str(weight_1m)
    table.rows[4].cells[0].text = t['6_months']
    table.rows[4].cells[1].text = str(weight_6m)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 5. REPORTES DE RECETAS
# ==========================================
def generate_recipes_report_pdf(analysis_data, profile_data, language="es"):
    t = get_translations(language)
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('MainTitle', parent=styles['Heading1'], fontSize=18, spaceAfter=20, textColor=colors.HexColor("#2c3e50"), alignment=1)
    subtitle_style = ParagraphStyle('SubTitle', parent=styles['Heading2'], fontSize=14, spaceAfter=10, textColor=colors.HexColor("#34495e"))
    normal_style = styles["Normal"]

    elements = []
    elements.append(Paragraph(f"{t['app_name']} - {t['recipes_report']}", title_style))
    elements.append(Paragraph(f"<b>{t['date_generated']}:</b> {datetime.now().strftime('%Y-%m-%d %H:%M')}", normal_style))
    elements.append(Paragraph(f"<b>{t['patient']}:</b> {profile_data.get('email', 'Usuario')}", normal_style))
    elements.append(Spacer(1, 20))

    nutrition_plan = analysis_data.get('health_plan', {}).get('nutrition', {})
    daily_calories = nutrition_plan.get('daily_calories', 0)
    protein = nutrition_plan.get('macronutrients', {}).get('protein', 0)

    elements.append(Paragraph(t['nutrition_summary'], subtitle_style))
    nutrition_data = [
        [t['metric'], t['value']],
        [t['daily_calories'], f"{daily_calories} kcal"],
        [t['protein'], f"{protein} g"]
    ]
    t_nutrition = Table(nutrition_data, colWidths=[200, 200])
    t_nutrition.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (1, 0), colors.HexColor("#34495e")),
        ('TEXTCOLOR', (0, 0), (1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor("#ecf0f1")),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    elements.append(t_nutrition)
    elements.append(Spacer(1, 20))

    elements.append(Paragraph(t['suggested_recipes'], subtitle_style))
    recipes_data = [
        [t['meal'], t['recipe'], t['protein_g']],
        [t['breakfast'], "Avena con frutas frescas, chía y un puñado de nueces" if language == "es" else "Oatmeal with fresh fruits, chia and a handful of nuts", str(round(protein * 0.2))],
        [t['lunch'], "Pechuga de pollo a la plancha con quinoa y mix de vegetales" if language == "es" else "Grilled chicken breast with quinoa and mixed vegetables", str(round(protein * 0.4))],
        [t['dinner'], "Ensalada ligera de salmón fresco, aguacate y espinaca" if language == "es" else "Light salad with fresh salmon, avocado and spinach", str(round(protein * 0.3))],
        [t['snack'], "Un yogur griego natural o un puñado de almendras tostadas" if language == "es" else "A natural Greek yogurt or a handful of roasted almonds", str(round(protein * 0.1))]
    ]
    t_recipes = Table(recipes_data, colWidths=[100, 250, 100])
    t_recipes.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#34495e")),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor("#ecf0f1")),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    elements.append(t_recipes)
    elements.append(Spacer(1, 20))

    doc.build(elements)
    buffer.seek(0)
    return buffer


def generate_recipes_report_excel(analysis_data, profile_data, language="es"):
    t = get_translations(language)
    buffer = io.BytesIO()
    workbook = xlsxwriter.Workbook(buffer, {'in_memory': True})
    header_format = workbook.add_format({'bold': True, 'bg_color': '#2c3e50', 'font_color': 'white', 'border': 1})

    ws1 = workbook.add_worksheet(t['nutrition_summary'])
    ws1.set_column('A:B', 25)
    ws1.write(0, 0, t['metric'], header_format)
    ws1.write(0, 1, t['value'], header_format)
    nutrition_plan = analysis_data.get('health_plan', {}).get('nutrition', {})
    ws1.write(1, 0, t['daily_calories'])
    ws1.write(1, 1, f"{nutrition_plan.get('daily_calories', 0)} kcal")
    ws1.write(2, 0, t['protein'])
    ws1.write(2, 1, f"{nutrition_plan.get('macronutrients', {}).get('protein', 0)} g")

    ws2 = workbook.add_worksheet(t['suggested_recipes'])
    ws2.set_column('A:C', 25)
    ws2.write_row(0, 0, [t['meal'], t['recipe'], t['protein_g']], header_format)
    protein = nutrition_plan.get('macronutrients', {}).get('protein', 0)
    recipes = [
        (t['breakfast'], "Avena con frutas frescas, chía y un puñado de nueces" if language == "es" else "Oatmeal with fresh fruits, chia and a handful of nuts", round(protein * 0.2)),
        (t['lunch'], "Pechuga de pollo a la plancha con quinoa y mix de vegetales" if language == "es" else "Grilled chicken breast with quinoa and mixed vegetables", round(protein * 0.4)),
        (t['dinner'], "Ensalada ligera de salmón fresco, aguacate y espinaca" if language == "es" else "Light salad with fresh salmon, avocado and spinach", round(protein * 0.3)),
        (t['snack'], "Un yogur griego natural o un puñado de almendras tostadas" if language == "es" else "A natural Greek yogurt or a handful of roasted almonds", round(protein * 0.1))
    ]
    for row_num, (momento, receta, prot) in enumerate(recipes, 1):
        ws2.write(row_num, 0, momento)
        ws2.write(row_num, 1, receta)
        ws2.write(row_num, 2, prot)

    workbook.close()
    buffer.seek(0)
    return buffer


def generate_recipes_report_word(analysis_data, profile_data, language="es"):
    t = get_translations(language)
    doc = Document()
    title = doc.add_heading(t['recipes_report'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.add_run(f"{t['date_generated']}: {datetime.now().strftime('%d/%m/%Y')}\n").bold = True
    p.add_run(t['patient_evaluation'])

    doc.add_heading(t['nutrition_summary_1'], level=1)
    nutrition_plan = analysis_data.get('health_plan', {}).get('nutrition', {})
    doc.add_paragraph(f"{t['daily_calories']}: {nutrition_plan.get('daily_calories', 0)} kcal")
    doc.add_paragraph(f"{t['protein']}: {nutrition_plan.get('macronutrients', {}).get('protein', 0)} g")

    doc.add_heading(t['suggested_recipes'], level=1)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = t['meal']
    hdr_cells[1].text = t['recipe']
    hdr_cells[2].text = t['protein_g']
    protein = nutrition_plan.get('macronutrients', {}).get('protein', 0)
    table.rows[1].cells[0].text = t['breakfast']
    table.rows[1].cells[1].text = "Avena con frutas frescas, chía y un puñado de nueces" if language == "es" else "Oatmeal with fresh fruits, chia and a handful of nuts"
    table.rows[1].cells[2].text = str(round(protein * 0.2))
    table.rows[2].cells[0].text = t['lunch']
    table.rows[2].cells[1].text = "Pechuga de pollo a la plancha con quinoa y mix de vegetales" if language == "es" else "Grilled chicken breast with quinoa and mixed vegetables"
    table.rows[2].cells[2].text = str(round(protein * 0.4))
    table.rows[3].cells[0].text = t['dinner']
    table.rows[3].cells[1].text = "Ensalada ligera de salmón fresco, aguacate y espinaca" if language == "es" else "Light salad with fresh salmon, avocado and spinach"
    table.rows[3].cells[2].text = str(round(protein * 0.3))
    table.rows[4].cells[0].text = t['snack']
    table.rows[4].cells[1].text = "Un yogur griego natural o un puñado de almendras tostadas" if language == "es" else "A natural Greek yogurt or a handful of roasted almonds"
    table.rows[4].cells[2].text = str(round(protein * 0.1))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 6. REPORTES DE EJERCICIO
# ==========================================
def generate_exercise_report_pdf(analysis_data, profile_data, language="es"):
    t = get_translations(language)
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('MainTitle', parent=styles['Heading1'], fontSize=18, spaceAfter=20, textColor=colors.HexColor("#2c3e50"), alignment=1)
    subtitle_style = ParagraphStyle('SubTitle', parent=styles['Heading2'], fontSize=14, spaceAfter=10, textColor=colors.HexColor("#34495e"))
    normal_style = styles["Normal"]

    elements = []
    elements.append(Paragraph(f"{t['app_name']} - {t['exercise_report']}", title_style))
    elements.append(Paragraph(f"<b>{t['date_generated']}:</b> {datetime.now().strftime('%Y-%m-%d %H:%M')}", normal_style))
    elements.append(Paragraph(f"<b>{t['patient']}:</b> {profile_data.get('email', 'Usuario')}", normal_style))
    elements.append(Spacer(1, 20))

    exercise_plan = analysis_data.get('health_plan', {}).get('exercise', {})
    elements.append(Paragraph(t['exercise_plan'], subtitle_style))
    exercise_data = [
        [t['type'], "Duración/Frecuencia" if language == "es" else "Duration/Frequency"],
        [t['cardio'], exercise_plan.get('cardio', 'N/A')],
        [t['strength'], exercise_plan.get('strength', 'N/A')],
        [t['flexibility'], exercise_plan.get('flexibility', 'N/A')]
    ]
    t_exercise = Table(exercise_data, colWidths=[200, 200])
    t_exercise.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (1, 0), colors.HexColor("#34495e")),
        ('TEXTCOLOR', (0, 0), (1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor("#ecf0f1")),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    elements.append(t_exercise)
    elements.append(Spacer(1, 20))

    doc.build(elements)
    buffer.seek(0)
    return buffer


def generate_exercise_report_excel(analysis_data, profile_data, language="es"):
    t = get_translations(language)
    buffer = io.BytesIO()
    workbook = xlsxwriter.Workbook(buffer, {'in_memory': True})
    header_format = workbook.add_format({'bold': True, 'bg_color': '#2c3e50', 'font_color': 'white', 'border': 1})

    ws1 = workbook.add_worksheet(t['exercise_plan'])
    ws1.set_column('A:B', 25)
    ws1.write(0, 0, t['type'], header_format)
    ws1.write(0, 1, "Duración/Frecuencia" if language == "es" else "Duration/Frequency", header_format)
    exercise_plan = analysis_data.get('health_plan', {}).get('exercise', {})
    exercises = [
        (t['cardio'], exercise_plan.get('cardio', 'N/A')),
        (t['strength'], exercise_plan.get('strength', 'N/A')),
        (t['flexibility'], exercise_plan.get('flexibility', 'N/A'))
    ]
    for row_num, (tipo, duracion) in enumerate(exercises, 1):
        ws1.write(row_num, 0, tipo)
        ws1.write(row_num, 1, duracion)

    workbook.close()
    buffer.seek(0)
    return buffer


def generate_exercise_report_word(analysis_data, profile_data, language="es"):
    t = get_translations(language)
    doc = Document()
    title = doc.add_heading(t['exercise_report'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.add_run(f"{t['date_generated']}: {datetime.now().strftime('%d/%m/%Y')}\n").bold = True
    p.add_run(t['patient_evaluation'])

    doc.add_heading(t['exercise_plan_1'], level=1)
    exercise_plan = analysis_data.get('health_plan', {}).get('exercise', {})
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = t['type']
    hdr_cells[1].text = "Duración/Frecuencia" if language == "es" else "Duration/Frequency"
    table.rows[1].cells[0].text = t['cardio']
    table.rows[1].cells[1].text = exercise_plan.get('cardio', 'N/A')
    table.rows[2].cells[0].text = t['strength']
    table.rows[2].cells[1].text = exercise_plan.get('strength', 'N/A')
    table.rows[3].cells[0].text = t['flexibility']
    table.rows[3].cells[1].text = exercise_plan.get('flexibility', 'N/A')

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
